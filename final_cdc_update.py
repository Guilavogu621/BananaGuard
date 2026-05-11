"""
Script pour refondre la section technique du Cahier des Charges.
- Section 3 ordonnée (Frontend -> Backend -> IA)
- Descriptions simplifiées (Style planning projet)
- Section IA complète (MobileNetV2, Dataset, Split, etc.)
- Police uniforme Arial
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from lxml import etree
import os

input_path = "Cahier des Charges — Version 1.0.1.docx"
output_path = "Cahier des Charges — Version 2.2.docx"

if not os.path.exists(input_path):
    print(f"Erreur: {input_path} introuvable.")
    exit(1)

doc = Document(input_path)

# Suppression de l'ancienne section 2.1 (elle sera intégrée en 3.3)
for p in doc.paragraphs:
    if p.text and "2.1 MODÈLE D'INTELLIGENCE ARTIFICIELLE" in p.text.upper():
        p._element.getparent().remove(p._element)
        break

# --- ÉTAPE 1: Insertion de la nouvelle Section 3 (Technologies) ---
body = doc.element.body

# Trouver l'endroit où insérer (avant "3. MÉTHODE")
target_element = None
for elem in body:
    if elem.tag == qn('w:p'):
        text = ''.join(t.text or '' for t in elem.iter(qn('w:t')))
        if "MÉTHODE DE DÉVELOPPEMENT" in text.upper():
            target_element = elem
            break

if target_element is not None:
    def create_p(text, bold=False, size=11, color=None, alignment=None):
        p = etree.Element(qn('w:p'))
        pPr = etree.SubElement(p, qn('w:pPr'))
        if alignment:
            jc = etree.SubElement(pPr, qn('w:jc'))
            jc.set(qn('w:val'), alignment)
        
        r = etree.SubElement(p, qn('w:r'))
        rPr = etree.SubElement(r, qn('w:rPr'))
        
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        
        if bold: etree.SubElement(rPr, qn('w:b'))
        if size:
            sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), str(size * 2))
        if color:
            c = etree.SubElement(rPr, qn('w:color'))
            c.set(qn('w:val'), color)
            
        t = etree.SubElement(r, qn('w:t'))
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        return p

    def create_table(headers, rows):
        tbl = etree.Element(qn('w:tbl'))
        tblPr = etree.SubElement(tbl, qn('w:tblPr'))
        tblStyle = etree.SubElement(tblPr, qn('w:tblStyle'))
        tblStyle.set(qn('w:val'), 'TableGrid')
        
        tr = etree.SubElement(tbl, qn('w:tr'))
        for h in headers:
            tc = etree.SubElement(tr, qn('w:tc'))
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            shd = etree.SubElement(tcPr, qn('w:shd'))
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1B5E20')
            p = etree.SubElement(tc, qn('w:p'))
            r = etree.SubElement(p, qn('w:r'))
            rPr = etree.SubElement(r, qn('w:rPr'))
            etree.SubElement(rPr, qn('w:b'))
            c = etree.SubElement(rPr, qn('w:color')); c.set(qn('w:val'), 'FFFFFF')
            t = etree.SubElement(r, qn('w:t')); t.text = h
            
        for row in rows:
            tr = etree.SubElement(tbl, qn('w:tr'))
            for val in row:
                tc = etree.SubElement(tr, qn('w:tc'))
                p = etree.SubElement(tc, qn('w:p'))
                r = etree.SubElement(p, qn('w:r'))
                t = etree.SubElement(r, qn('w:t')); t.text = val
        return tbl

    new_elems = [
        create_p("3. ARCHITECTURE TECHNIQUE ET TECHNOLOGIES", bold=True, size=14, color="1B5E20"),
        create_p("Sélection des technologies pour garantir la performance et la scalabilité du projet.", size=10),
        
        create_p("3.1 Technologies Frontend", bold=True, size=12, color="388E3C"),
        create_table(["Technologie", "Rôle"], [
            ["React 19 & Vite", "Développement de l'interface utilisateur."],
            ["React Router DOM", "Gestion de la navigation."],
            ["Axios", "Appels API vers le backend."],
            ["Framer Motion", "Animations et transitions."],
            ["Lucide React", "Bibliothèque d'icônes."],
            ["Vite PWA Plugin", "Support mobile et mode hors-ligne."],
        ]),
        
        create_p("3.2 Technologies Backend", bold=True, size=12, color="388E3C"),
        create_table(["Technologie", "Rôle"], [
            ["FastAPI & Uvicorn", "Serveur API asynchrone."],
            ["SQLAlchemy & SQLite", "Gestion de la base de données."],
            ["Pydantic", "Validation des schémas de données."],
            ["JWT (python-jose)", "Sécurisation de l'authentification."],
            ["Bcrypt", "Hachage des mots de passe."],
            ["Pandas", "Export des données en CSV."],
        ]),
        
        create_p("3.3 Spécifications de l'IA", bold=True, size=12, color="388E3C"),
        create_table(["Paramètre", "Détail"], [
            ["Modèle", "MobileNetV2 (Transfer Learning — ImageNet)"],
            ["Framework", "TensorFlow / Keras 2.19"],
            ["Dataset total", "1 600 images, 4 classes, 400 images/classe"],
            ["Sources", "Kaggle BananaLSD + Banana Disease Recognition"],
            ["Augmentation", "Rotation, zoom, flip horizontal, luminosité"],
            ["Split", "70% train / 15% validation / 15% test"],
        ]),
        create_p("", size=10)
    ]

    for elem in reversed(new_elems):
        target_element.addprevious(elem)

# --- ÉTAPE 2: Mise à jour Jira et Agile ---
for table in doc.tables:
    try:
        first_cell_text = table.cell(0, 0).text
        if "Outil" in first_cell_text:
            new_row = table.add_row()
            new_row.cells[0].text = "Jira"
            new_row.cells[1].text = "Gestion de projet Agile : suivi des tâches et sprints."
    except:
        continue

# --- ÉTAPE 3: Renumérotation et Police ---
for p in doc.paragraphs:
    if p.text is None: continue
    txt = p.text.strip()
    if txt.startswith("3. MÉTHODE"): p.text = p.text.replace("3.", "4.", 1)
    elif txt.startswith("4. PÉRIMÈTRE"): p.text = p.text.replace("4.", "5.", 1)
    for run in p.runs:
        run.font.name = 'Arial'

# Mise à jour Version et Date
for p in doc.paragraphs:
    if p.text is None: continue
    if "Version 1.0.1" in p.text: p.text = p.text.replace("Version 1.0.1", "Version 2.2")
    if "16 Avril" in p.text: p.text = p.text.replace("16 Avril 2026", "29 Avril 2026")

doc.save(output_path)
print(f"✅ Document sauvegardé : {output_path}")
