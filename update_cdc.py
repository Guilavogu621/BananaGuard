"""
Script pour mettre à jour le Cahier des Charges BananaGuard.
Ajoute : Technologies Frontend/Backend + Jira comme outil de gestion.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import copy

doc = Document("Cahier des Charges — Version 1.0.1.docx")

# --- HELPER FUNCTIONS ---
def make_bold_run(paragraph, text, size=11, color=None):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def make_run(paragraph, text, size=11, color=None):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_styled_table(doc, headers, rows, insert_before=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        make_bold_run(p, h, size=10, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '1B5E20'
        })
        shading.append(shd)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            make_run(p, val, size=10)
    return table

# --- FIND THE RIGHT INSERTION POINT ---
# We need to find where section "3. MÉTHODE DE DÉVELOPPEMENT" is
# and insert the technologies section BEFORE it (as section 3)
# Then the existing section 3 becomes section 4, and section 4 becomes 5

paragraphs = doc.paragraphs
insert_index = None
tools_table_index = None

for i, p in enumerate(paragraphs):
    text = p.text.strip()
    # Find "3. MÉTHODE DE DÉVELOPPEMENT"
    if "MÉTHODE DE DÉVELOPPEMENT" in text or "METHODE DE DEVELOPPEMENT" in text:
        insert_index = i
    # Find "Outils Agile" to locate the tools table
    if "Outils Agile" in text or "outils Agile" in text:
        tools_table_index = i

print(f"Section 'Méthode de développement' trouvée à l'index: {insert_index}")
print(f"Section 'Outils Agile' trouvée à l'index: {tools_table_index}")

# --- STEP 1: Update the existing tools table to add Jira ---
# Find the table that comes after "Outils Agile utilisés"
for i, table in enumerate(doc.tables):
    first_cell = table.rows[0].cells[0].text.strip()
    if "Outil" in first_cell:
        print(f"Table 'Outils' trouvée (table index {i})")
        # Add a new row for Jira at the top of data rows
        new_row = table.add_row()
        new_row.cells[0].text = "Jira"
        new_row.cells[1].text = "Gestion de projet Agile : backlog, sprints, user stories, suivi des tâches, réunions (PV), rapports d'avancement"
        print("✅ Jira ajouté à la table des outils Agile")
        break

# --- STEP 2: Add Technologies section ---
# We'll insert new paragraphs before "3. MÉTHODE DE DÉVELOPPEMENT"
# First, let's work with the XML directly for precise insertion

from docx.oxml.ns import qn
from lxml import etree

body = doc.element.body

# Find the paragraph element for section 3
target_element = None
for elem in body:
    if elem.tag == qn('w:p'):
        text = ''.join(t.text or '' for t in elem.iter(qn('w:t')))
        if "MÉTHODE DE DÉVELOPPEMENT" in text or "METHODE DE DEVELOPPEMENT" in text:
            target_element = elem
            break

if target_element is not None:
    print("Insertion de la section Technologies...")

    def create_paragraph(text, bold=False, size=11, color=None, heading_level=None, alignment=None):
        """Create a w:p element with styled text."""
        p = etree.SubElement(body, qn('w:p'))  # temp, will be moved
        body.remove(p)  # remove from body, we'll insert manually

        pPr = etree.SubElement(p, qn('w:pPr'))
        if heading_level:
            pStyle = etree.SubElement(pPr, qn('w:pStyle'))
            pStyle.set(qn('w:val'), f'Heading{heading_level}')
        if alignment:
            jc = etree.SubElement(pPr, qn('w:jc'))
            jc.set(qn('w:val'), alignment)

        r = etree.SubElement(p, qn('w:r'))
        rPr = etree.SubElement(r, qn('w:rPr'))
        if bold:
            etree.SubElement(rPr, qn('w:b'))
        if size:
            sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), str(size * 2))
            szCs = etree.SubElement(rPr, qn('w:szCs'))
            szCs.set(qn('w:val'), str(size * 2))
        if color:
            c = etree.SubElement(rPr, qn('w:color'))
            c.set(qn('w:val'), color)

        t = etree.SubElement(r, qn('w:t'))
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        return p

    def create_table(headers, rows):
        """Create a w:tbl element."""
        tbl = etree.SubElement(body, qn('w:tbl'))
        body.remove(tbl)

        # Table properties
        tblPr = etree.SubElement(tbl, qn('w:tblPr'))
        tblStyle = etree.SubElement(tblPr, qn('w:tblStyle'))
        tblStyle.set(qn('w:val'), 'TableGrid')
        tblW = etree.SubElement(tblPr, qn('w:tblW'))
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblLook = etree.SubElement(tblPr, qn('w:tblLook'))
        tblLook.set(qn('w:val'), '04A0')

        # Table grid
        tblGrid = etree.SubElement(tbl, qn('w:tblGrid'))
        for _ in headers:
            etree.SubElement(tblGrid, qn('w:gridCol'))

        def make_cell(text, is_header=False):
            tc = etree.Element(qn('w:tc'))
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            if is_header:
                shd = etree.SubElement(tcPr, qn('w:shd'))
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '1B5E20')
            p = etree.SubElement(tc, qn('w:p'))
            r = etree.SubElement(p, qn('w:r'))
            rPr = etree.SubElement(r, qn('w:rPr'))
            if is_header:
                etree.SubElement(rPr, qn('w:b'))
                c = etree.SubElement(rPr, qn('w:color'))
                c.set(qn('w:val'), 'FFFFFF')
            sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), '20')
            t_elem = etree.SubElement(r, qn('w:t'))
            t_elem.text = text
            t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            return tc

        # Header row
        header_row = etree.SubElement(tbl, qn('w:tr'))
        for h in headers:
            header_row.append(make_cell(h, is_header=True))

        # Data rows
        for row_data in rows:
            tr = etree.SubElement(tbl, qn('w:tr'))
            for val in row_data:
                tr.append(make_cell(val))

        return tbl

    # Elements to insert (in order, BEFORE the target)
    elements_to_insert = []

    # Section title
    elements_to_insert.append(
        create_paragraph("3. ARCHITECTURE TECHNIQUE ET TECHNOLOGIES", bold=True, size=14, color="1B5E20")
    )
    elements_to_insert.append(
        create_paragraph("Cette section détaille l'ensemble des technologies, frameworks et bibliothèques utilisés pour le développement de BananaGuard.", size=11)
    )

    # 3.1 Frontend
    elements_to_insert.append(
        create_paragraph("3.1 Technologies Frontend", bold=True, size=12, color="388E3C")
    )

    frontend_table = create_table(
        ["Technologie", "Version", "Rôle"],
        [
            ["React", "19.2", "Bibliothèque UI — construction de l'interface utilisateur en composants réactifs"],
            ["Vite", "8.0", "Outil de build et serveur de développement ultra-rapide (remplacement de Webpack)"],
            ["React Router DOM", "7.14", "Gestion du routage SPA — navigation entre les pages sans rechargement"],
            ["Axios", "1.15", "Client HTTP pour les appels API vers le backend FastAPI"],
            ["Framer Motion", "12.38", "Bibliothèque d'animations et de transitions fluides pour l'interface"],
            ["Lucide React", "1.11", "Pack d'icônes SVG modernes et légers pour l'iconographie"],
            ["CSS3 (Vanilla)", "—", "Stylisation complète de l'interface sans framework CSS externe"],
            ["Vite PWA Plugin", "1.2", "Progressive Web App — accès hors-ligne et installation sur mobile"],
            ["ESLint", "10.2", "Linter JavaScript pour le contrôle qualité du code"],
        ]
    )
    elements_to_insert.append(frontend_table)

    # 3.2 Backend
    elements_to_insert.append(
        create_paragraph("3.2 Technologies Backend", bold=True, size=12, color="388E3C")
    )

    backend_table = create_table(
        ["Technologie", "Version", "Rôle"],
        [
            ["Python", "3.11", "Langage de programmation principal du serveur"],
            ["FastAPI", "Dernière", "Framework web asynchrone haute performance pour l'API REST"],
            ["Uvicorn", "Standard", "Serveur ASGI pour exécuter l'application FastAPI"],
            ["SQLAlchemy", "Dernière", "ORM — mapping objet-relationnel pour la base de données"],
            ["SQLite", "3", "Base de données embarquée légère (fichier bananaguard.db)"],
            ["Pydantic", "v2", "Validation des données et sérialisation des schémas API"],
            ["JWT (python-jose)", "Dernière", "Authentification par tokens JSON Web Token (HS256)"],
            ["Passlib + Bcrypt", "4.0", "Hachage sécurisé des mots de passe utilisateur"],
            ["TensorFlow / Keras", "CPU", "Exécution du modèle IA MobileNetV2 pour la prédiction"],
            ["Pillow (PIL)", "Dernière", "Traitement et redimensionnement des images uploadées"],
            ["Pandas", "Dernière", "Export des historiques au format CSV"],
        ]
    )
    elements_to_insert.append(backend_table)

    # 3.3 IA
    elements_to_insert.append(
        create_paragraph("3.3 Intelligence Artificielle", bold=True, size=12, color="388E3C")
    )

    ia_table = create_table(
        ["Technologie", "Rôle"],
        [
            ["TensorFlow / Keras", "Framework d'entraînement et d'inférence du modèle CNN"],
            ["MobileNetV2 (Transfer Learning)", "Architecture du réseau de neurones — pré-entraîné sur ImageNet"],
            ["Google Colab", "Environnement cloud pour l'entraînement du modèle (GPU gratuit)"],
            ["Jupyter Notebook", "Documentation et expérimentation du pipeline d'entraînement"],
            ["NumPy", "Manipulation des tableaux de données et des prédictions"],
        ]
    )
    elements_to_insert.append(ia_table)

    # Empty paragraph separator
    elements_to_insert.append(create_paragraph("", size=11))

    # Insert all elements before the target (section 3 → becomes section 4)
    for elem in elements_to_insert:
        target_element.addprevious(elem)

    print("✅ Section 'Architecture Technique et Technologies' insérée")

# --- STEP 3: Renumber existing sections ---
# Section 3 → 4, Section 4 → 5
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith("3. MÉTHODE") or text.startswith("3. METHODE"):
        for run in p.runs:
            run.text = run.text.replace("3.", "4.", 1)
    elif text.startswith("3.") and "Outils Agile" in text:
        for run in p.runs:
            run.text = run.text.replace("3.", "4.", 1)
    elif text.startswith("4. PÉRIMÈTRE") or text.startswith("4. PERIMETRE"):
        for run in p.runs:
            run.text = run.text.replace("4.", "5.", 1)

# Also update the version in the footer-like paragraph
for p in doc.paragraphs:
    if "Version 1.2" in p.text:
        for run in p.runs:
            run.text = run.text.replace("Version 1.2", "Version 2.0")
    if "20 Avril 2026" in p.text and "Mis à jour" in p.text:
        for run in p.runs:
            run.text = run.text.replace("20 Avril 2026", "29 Avril 2026")

# Save the updated document
output_path = "Cahier des Charges — Version 2.0.docx"
doc.save(output_path)
print(f"\n✅ Document sauvegardé : {output_path}")
