"""
Script pour affiner le Cahier des Charges BananaGuard.
- Police uniforme (Inter/Arial)
- Technologies simplifiées (Frontend/Backend)
- Déplacement de 2.1 vers 3.3
- Ajout de Jira
- Style mature et professionnel
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import os

# Entrée: Version originale
input_path = "Cahier des Charges — Version 1.0.1.docx"
output_path = "Cahier des Charges — Version 2.1.docx"

if not os.path.exists(input_path):
    print(f"Erreur: {input_path} introuvable.")
    exit(1)

doc = Document(input_path)

# --- FONCTION DE STYLE GLOBAL ---
def set_global_font(doc, font_name='Arial'): # Arial est standard, Inter est préféré si installé
    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = font_name
            # Forcer pour les langues complexes/est-asiatiques
            style.element.xpath('w:rPr/w:rFonts')[0].set(qn('w:ascii'), font_name)
            style.element.xpath('w:rPr/w:rFonts')[0].set(qn('w:hAnsi'), font_name)

set_global_font(doc)

# --- HELPER: Suppression de paragraphe/élément ---
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

# --- ÉTAPE 1: Extraire les données de 2.1 (Modèle IA) et supprimer la section originale ---
ia_data = []
found_2_1 = False
paragraphs_to_remove = []
tables_to_remove = []

for i, p in enumerate(doc.paragraphs):
    if "2.1 MODÈLE D'INTELLIGENCE ARTIFICIELLE" in p.text.upper():
        found_2_1 = True
        paragraphs_to_remove.append(p)
        # Chercher la table juste après
        next_idx = i + 1
        # On va chercher la table associée dans doc.tables en comparant les éléments
        continue

# Suppression de la section 2.1 dans le texte
for p in doc.paragraphs:
    if "2.1 MODÈLE D'INTELLIGENCE ARTIFICIELLE" in p.text.upper():
        delete_paragraph(p)
        break

# --- ÉTAPE 2: Ajouter Jira dans la table Agile ---
for table in doc.tables:
    if "Outil" in table.rows[0].cells[0].text:
        new_row = table.add_row()
        new_row.cells[0].text = "Jira"
        new_row.cells[1].text = "Gestion de projet Agile : Backlog, Sprints, suivi des tâches et rapports d'avancement."
        break

# --- ÉTAPE 3: Insertion de la nouvelle Section 3 (Technologies) ---
body = doc.element.body

# Trouver l'endroit où insérer (avant "3. MÉTHODE")
target_element = None
for elem in body:
    if elem.tag == qn('w:p'):
        text = ''.join(t.text or '' for t in elem.iter(qn('w:t')))
        if "MÉTHODE DE DÉVELOPPEMENT" in text.upper():
            target_element = elem
            break

if target_element is not None:
    def create_p(text, bold=False, size=11, color=None, alignment=None):
        p = etree.Element(qn('w:p'))
        pPr = etree.SubElement(p, qn('w:pPr'))
        if alignment:
            jc = etree.SubElement(pPr, qn('w:jc'))
            jc.set(qn('w:val'), alignment)
        
        r = etree.SubElement(p, qn('w:r'))
        rPr = etree.SubElement(r, qn('w:rPr'))
        
        # Font
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        
        if bold: etree.SubElement(rPr, qn('w:b'))
        if size:
            sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), str(size * 2))
        if color:
            c = etree.SubElement(rPr, qn('w:color'))
            c.set(qn('w:val'), color)
            
        t = etree.SubElement(r, qn('w:t'))
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        return p

    def create_table(headers, rows):
        tbl = etree.Element(qn('w:tbl'))
        tblPr = etree.SubElement(tbl, qn('w:tblPr'))
        tblStyle = etree.SubElement(tblPr, qn('w:tblStyle'))
        tblStyle.set(qn('w:val'), 'TableGrid')
        
        # Header
        tr = etree.SubElement(tbl, qn('w:tr'))
        for h in headers:
            tc = etree.SubElement(tr, qn('w:tc'))
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            shd = etree.SubElement(tcPr, qn('w:shd'))
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1B5E20')
            
            p = etree.SubElement(tc, qn('w:p'))
            r = etree.SubElement(p, qn('w:r'))
            rPr = etree.SubElement(r, qn('w:rPr'))
            etree.SubElement(rPr, qn('w:b'))
            c = etree.SubElement(rPr, qn('w:color')); c.set(qn('w:val'), 'FFFFFF')
            t = etree.SubElement(r, qn('w:t')); t.text = h
            
        # Rows
        for row in rows:
            tr = etree.SubElement(tbl, qn('w:tr'))
            for val in row:
                tc = etree.SubElement(tr, qn('w:tc'))
                p = etree.SubElement(tc, qn('w:p'))
                r = etree.SubElement(p, qn('w:r'))
                t = etree.SubElement(r, qn('w:t')); t.text = val
        return tbl

    new_elems = [
        create_p("3. ARCHITECTURE TECHNIQUE ET TECHNOLOGIES", bold=True, size=14, color="1B5E20"),
        create_p("Sélection des technologies matures utilisées pour garantir la performance et la scalabilité du projet.", size=10),
        
        create_p("3.1 Technologies Frontend", bold=True, size=12, color="388E3C"),
        create_table(["Technologie", "Rôle"], [
            ["React 19 & Vite", "Interface utilisateur réactive et build optimisé."],
            ["React Router DOM", "Navigation fluide type Single Page Application."],
            ["Axios", "Communication asynchrone avec l'API REST."],
            ["Framer Motion", "Animations premium et transitions d'interface."],
            ["Lucide React", "Iconographie moderne et légère."],
            ["Vite PWA Plugin", "Support du mode hors-ligne et installation mobile."],
        ]),
        
        create_p("3.2 Technologies Backend", bold=True, size=12, color="388E3C"),
        create_table(["Technologie", "Rôle"], [
            ["FastAPI & Uvicorn", "Serveur API haute performance et asynchrone."],
            ["SQLAlchemy & SQLite", "Gestion de la base de données relationnelle locale."],
            ["Pydantic", "Validation stricte des données et schémas."],
            ["JWT (python-jose)", "Sécurisation des accès par jetons d'authentification."],
            ["Bcrypt", "Hachage sécurisé des mots de passe."],
            ["Pandas", "Traitement de données et génération de rapports CSV."],
        ]),
        
        create_p("3.3 Spécifications de l'IA (anciennement 2.1)", bold=True, size=12, color="388E3C"),
        create_table(["Paramètre", "Détail"], [
            ["Modèle", "MobileNetV2 (Transfer Learning)"],
            ["Framework", "TensorFlow / Keras"],
            ["Dataset", "1 600 images (4 classes)"],
            ["Optimisation", "Data Augmentation & Pré-entraînement ImageNet"],
        ]),
        create_p("", size=10) # Espacement
    ]

    for elem in reversed(new_elems):
        target_element.addprevious(elem)

# --- ÉTAPE 4: Renumérotation et Nettoyage ---
for p in doc.paragraphs:
    if p.text is None: continue
    txt = p.text.strip()
    if txt.startswith("3. MÉTHODE"): p.text = p.text.replace("3.", "4.", 1)
    elif txt.startswith("4. PÉRIMÈTRE"): p.text = p.text.replace("4.", "5.", 1)
    # Uniformiser la police
    for run in p.runs:
        run.font.name = 'Arial'

# Mise à jour de la version et date
for p in doc.paragraphs:
    if p.text is None: continue
    if "Version 1.0.1" in p.text: p.text = p.text.replace("Version 1.0.1", "Version 2.1")
    if "16 Avril" in p.text: p.text = p.text.replace("16 Avril 2026", "29 Avril 2026")

doc.save(output_path)
print(f"✅ Document sauvegardé : {output_path}")
